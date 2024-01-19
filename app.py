#For Streamlit libraries
import pandas as pd 
import streamlit as st
from datetime import datetime, timezone, timedelta
#For Twilio
import time
import os
import twilio
from twilio.rest import Client
import base64
import requests
from requests.structures import CaseInsensitiveDict
from io import BytesIO
from docx import Document
from docx.shared import Inches
import speech_recognition as sr #add this library
import urllib.request

list_audioLink1 = ["https://storage.googleapis.com/us.artifacts.kaga-shi-gomi-chatbot-ngdm.appspot.com/twilio-audio/new_voice/MZ1.mp3","https://storage.googleapis.com/us.artifacts.kaga-shi-gomi-chatbot-ngdm.appspot.com/twilio-audio/new_voice/MZ2.mp3"]
list_audioLink2 = ["https://storage.googleapis.com/us.artifacts.kaga-shi-gomi-chatbot-ngdm.appspot.com/twilio-audio/aug15/sugimoto_sam_de_irasahai.mp3","https://storage.googleapis.com/us.artifacts.kaga-shi-gomi-chatbot-ngdm.appspot.com/twilio-audio/aug15/ishou_de_irashaimasu_ka.mp3"]
phase2_voice = ["https://storage.googleapis.com/us.artifacts.kaga-shi-gomi-chatbot-ngdm.appspot.com/twilio-audio/AG1.mp3","https://storage.googleapis.com/us.artifacts.kaga-shi-gomi-chatbot-ngdm.appspot.com/twilio-audio/AG2.mp3"]
senjitu_changer = ["2022年10月","2022年9月"]


#Setting Page title and header
st.set_page_config(page_title="Phase_3 Debt Collection project")
st.header("債権回収自動コールシステム_PHASE-3")
st.subheader("下記のところにエクセルのファイルアップしてください")





document = Document()
# Add a main heading to the document
document.add_heading("通話履歴レポートのテキストの要約", level=0)

def add_logs(str1,str2):


    # Add a subheading to the document
    document.add_heading(str1, level=1)
    document.add_paragraph(str2)


#Place for uploading excel file

upload_file = st.file_uploader("xlsxファイルをここにアップしてください", type="xlsx")
if upload_file:
    st.markdown("正しいファイルタイプ")
    df = pd.read_excel(upload_file, engine="openpyxl")
    show_df = df.iloc[: , :] #df.iloc[: , :-2]
    # date_added = ["2022年10月","2022年9月"]
    # show_df['Date_Added'] = date_added
    show_df.rename(columns = {'Sr.':'番号', 'Name':'氏名', 'TEL':'電話番号', 'thing':'商品・サービス', 'amount':'購入金額', 'Date_Added':'購入日'}, inplace = True)
    hide_table_row_index = """
            <style>
            thead tr th:first-child {display:none}
            tbody th {display:none}
            </style>
            """
    st.markdown(hide_table_row_index, unsafe_allow_html=True)
    
    st.table(show_df)
   
    st.markdown("今回は電話かけるのは " + str(len(df))+"名様になります")


    
    if st.button("実行"):
        #wordings = fun_twilio



    
        
        name_list = df["Name"].tolist()
        mobile_numbers_list1 = df["TEL"].tolist()
        audioLink1= list_audioLink1 #df["audio_link"].tolist()
        audioLink2= list_audioLink2 #df["audio_link2"].tolist()
        item_name=df["thing"].tolist()
        money_left=df["amount"].tolist()
        senjitu_changer=df["Date_Added"].tolist()

        print("I am here man")
        print(name_list)
        print(mobile_numbers_list1)
        print(type(mobile_numbers_list1[0]))

        mobile_numbers_list=[]
        for i in range(0,len(mobile_numbers_list1)):
            add_this= "+81"+str(mobile_numbers_list1[i])
            mobile_numbers_list.append(add_this)
            print(mobile_numbers_list[i])
        
        print(mobile_numbers_list)
        
    

        
        #making twilio programme work in jikkou button
        p=0
        for k in range (0,len(mobile_numbers_list)):

            #denwabango in mobile_numbers_list:
            

            account_sid = st.secrets["account_sid"]
            auth_token = st.secrets["auth_token"]
            client = Client(account_sid, auth_token)
            good_number = "0" + str(mobile_numbers_list[k][3:13])

            
            number1 = good_number[0:3]+"x"+good_number[3:7]+"x"+good_number[7:11]
            numbers_break = list(number1)
            print(numbers_break)


            test_string = {"0" : "せろ", "1" : "いち","2" : "に","3" : "さん","4" : "よん","5" : "ごう","6" : "ろく","7" : "なな","8" : "はち","9" : "きゅう","x" : "の、"}

            empty_jap_eng = ""
            for i in numbers_break:
                empty_jap_eng = empty_jap_eng + test_string[i]
            print(empty_jap_eng)

            person_called = name_list[p]+"様"
            execution = client.studio \
                              .flows(st.secrets["flow"]) \
                              .executions \
                              .create(parameters={
                                'userName' : name_list[p]+"様",
                                'name_checker' : name_list[p][0:2],
                                'audioLink' : audioLink1[1],
                                'audioLink2' : audioLink2[1],
                                'item_to_twilio' : item_name[k],
                                'amount_to_twilio' : money_left[k],
                                'phase2_voice_to_twilio' : phase2_voice[1],
                                'senjitu_changer_to_twilio' : senjitu_changer[k],
                                'mobile_number' : empty_jap_eng,
                                },to=mobile_numbers_list[k], from_=st.secrets["twilio_number"]).fetch()


            print(execution.sid)

            
            calls = client.calls.list(to=mobile_numbers_list[k], limit=1)
            for record in calls:
                sid_to_be_used = record.sid
                utc_time = record.date_created
                date_called1 = utc_time.astimezone(timezone(timedelta(hours=9)))
                date_called = date_called1.strftime("%m/%d/%Y_%H:%M:%S")
                dare_called_for_logs = date_called1.strftime("%m/%d/%Y")
                #st.write(type(date_called))
                #st.write(record.sid)

            col1, col2, col3,col4 = st.columns(4)
            
            col2 = st.empty()
            

            col2.markdown("🟢電話中")

            with col4:
                option = st.selectbox(label = "お選びください",options =('ストップ', 'プレ'),key=k)
                st.write('You selected:', option)
            
            with col3:
                if option == "ストップ": 
                    st.write("🟠RESUME")
                else:
                    st.write("🔴終了")
                
            with col1:
                st.write(name_list[k])
                st.write(mobile_numbers_list[k])

            

            #Getting input after call
                
            user_input = st.text_area("コメント", key = k,placeholder="ここにコメント入力ください")
            st.markdown(user_input)


            #people = "World is haaaa"
            ## this below time is lag between first statment between users
            time.sleep(10)
      

             ## Checking first call has ended or not 
            kya_ho_raha_he = client.studio \
                                    .flows(st.secrets["flow"]) \
                                    .executions(execution.sid) \
                                    .fetch()
                                    
            print(kya_ho_raha_he.status)

            if (kya_ho_raha_he.status == "ended"):
                col2.markdown("🔴終了")

            #for i in range(0,len(name_list)):
                
            
            while(kya_ho_raha_he.status == "active"):
                kya_ho_raha_he = client.studio \
                                       .flows(st.secrets["flow"]) \
                                       .executions(execution.sid) \
                                       .fetch()
                # every 6.2 seconds fetch the information about what is happening
                time.sleep(3.2)

            print("###############################################################")
            print("Call has ended")
            
            #Getting recording sid
            url_set = f"https://api.twilio.com/2010-04-01/Accounts/{account_sid}/Calls/{sid_to_be_used}/Recordings.json"
            url1 = url_set
            headers = CaseInsensitiveDict()
            headers["Authorization"] = st.secrets["auth_one"]
            resp = requests.get(url1, headers=headers)
            resp_to_json = resp.json()
            #st.write(resp_to_json["recordings"][0]["sid"])
            recording_sid = resp_to_json["recordings"][0]["sid"]
            
            #Fetching the recording
            url = f"https://api.twilio.com/2010-04-01/Accounts/{account_sid}/Recordings/{recording_sid}.mp3"
            filename = f"{person_called}_{mobile_numbers_list[k]}_{date_called}_.mp3"
            filename_for_log = f"{person_called}_{mobile_numbers_list[k]}_{date_called}"


            response = requests.get(url)
            file_bytes = response.content
            st.markdown(f"{person_called}_{mobile_numbers_list[k]}の録音保存: <a href='data:audio/mp3;base64,{base64.b64encode(file_bytes).decode()}' download='{filename}'>ダウンロード</a>", unsafe_allow_html=True)


            #Adding logs functions for calls
            r = sr.Recognizer()

            # load the Japanese audio file
            japanese_audio_file = sr.AudioFile("japanese_file.wav")
            audio_url = f"https://api.twilio.com/2010-04-01/Accounts/{account_sid}/Recordings/{recording_sid}.wav"

            # download the audio file and save it locally
            urllib.request.urlretrieve(audio_url, "japanese_file.wav")

            # read the Japanese audio data from the file
            with japanese_audio_file as source:
                japanese_audio_data = r.record(source)
            # convert the Japanese audio data to text
            result2_logs = r.recognize_google(japanese_audio_data, language="ja-JP", show_all=True)
            japanese_text = result2_logs['alternative'][0]['transcript']


            add_logs(filename_for_log, japanese_text)


            # Save the document to a BytesIO buffer


            
            if (kya_ho_raha_he.status == "ended"):
                col2.markdown("🔴終了")

            st.write("----------------------------------------------------------------------")
            ### This is after one call ends time check here what to do
            #time.sleep(6)
            time.sleep(4)
 
            #time.sleep(5)
            p = p+1
 
    
        # Save the document to a BytesIO buffer
        doc_buffer = BytesIO()
        document.save(doc_buffer)
        doc_buffer.seek(0)        
        href = f'テキストログレポートを <a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{base64.b64encode(doc_buffer.getvalue()).decode()}" download="dare_called_for_logs.docx">ダウンロード</a>'
        st.markdown(href, unsafe_allow_html=True)

            #Makingthe web page having details about programme
    

   
